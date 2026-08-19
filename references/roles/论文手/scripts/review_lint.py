# -*- coding: utf-8 -*-
"""终审确定性 lint（只读，不修改任何交付物）。

把论文终审清单 A 类（C1/C5/C6）可机械核验的检查脚本化；B 类（三档逐字支撑、
反向锚点、定性结论数字锚点）留给 W2 质检 Subagent 判定，本脚本不替代 W2。

内置检查模式（对应终审清单.md A 类）：
  C6a 占位符/编辑残留    待补/TODO/占位/待定/TBD/（待/【 …
  C6b 双重编号           LaTeX: \\caption{表/图 N…} 手写编号撞自动编号；
                          Word: 正文「图 N/表 N」重号
  C1  孤儿数字           正文叙述小数在 results/ 无同源（按相对容差，可用白名单剔除）
  C1  口径锚点           论文关键声称值 vs results 真值（config 驱动）
  C5  表述↔实现          声称的算法行为 vs 实现/结果（config 驱动）
  C1① 判据方向           病态/适定断言 vs 条件数结果文件（config 驱动）

项目特异性（锚点值、声称规则、判据结果文件、孤儿数字白名单）放
<PROJECT_ROOT>/review_lint.json；脚本本体内置在 skill，全项目共用。

用法：
  python review_lint.py --paper <main.tex|完整论文.docx|tex项目目录> --results <results目录>
        [--config review_lint.json] [--demo] [--relax-orphans] [--tol 0.02]

退出码：
  正常模式  0 = 无命中且每个配置检查（anchors/claims/judgements）至少命中 1 处；
            1 = 有命中，或存在 0 命中的配置项（配置没对上论文 → FAIL，不允许进 W2）
  --demo    0 = 全部注入缺陷被捕获；1 = 存在未被捕获的注入缺陷
"""
import re, json, glob, sys, os, argparse

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

# ---------------- 文本抽取 ----------------
def load_text(paper):
    """返回 (文本, source_type)，source_type ∈ {'tex','docx'}。doc.paragraphs 不含表格文本→数据表自动排除。"""
    if os.path.isdir(paper):
        parts = []
        for f in sorted(glob.glob(os.path.join(paper, '**', '*.tex'), recursive=True)):
            parts.append(open(f, encoding='utf-8', errors='replace').read())
        return '\n'.join(parts), 'tex'
    if str(paper).lower().endswith('.docx'):
        try:
            from docx import Document
        except ImportError:
            sys.exit('缺少 python-docx，无法读取 Word 论文：pip install python-docx')
        doc = Document(paper)
        return '\n'.join(p.text for p in doc.paragraphs if p.text.strip()), 'docx'
    return open(paper, encoding='utf-8', errors='replace').read(), 'tex'


def _strip_tables(s):
    # 用等量换行占位，保证剥表后后续段落的 L# 不变
    def _repl(m):
        return '\n' * m.group(0).count('\n')
    return re.sub(r'\\begin\{(tabular|tabularx|table)\}.*?\\end\{(tabular|tabularx|table)\}',
                  _repl, s, flags=re.S)


def _strip_layout_args(s):
    # 用等量换行占位，保证剥掉 includegraphics 参数后行号不变
    def _repl(m):
        return '\n' * m.group(0).count('\n')
    return re.sub(r'\\includegraphics\[[^]]*\]\{[^}]*\}', _repl, s)


def _nums(s):
    out = set()
    for m in re.finditer(r'(?<![A-Za-z_])\d+\.\d+(?:[eE][+-]?\d+)?', s):
        out.add(float(m.group()))
    return out


def _near(v, res_nums, tol):
    return any(abs(v - r) <= max(tol * abs(r), 0.01) for r in res_nums)


def _load_res_nums(results_dir):
    t = ''
    for pat in ('*.json', '*.csv'):
        for f in glob.glob(os.path.join(results_dir, pat)):
            try:
                t += open(f, encoding='utf-8', errors='replace').read()
            except OSError:
                pass
    return _nums(t)


def _ln(text, m):
    return text[:m.start()].count('\n') + 1


def _plain(text):
    # 供 C1/C5/C1① 检查共用：剥数据表与 includegraphics 参数，只对正文叙述文本做匹配
    return _strip_layout_args(_strip_tables(text))


def _ctx(text, m, w=18):
    return text[max(0, m.start() - w):m.end() + w].replace('\n', ' ')


# ---------------- 配置 ----------------
def load_config(cfg_path):
    if cfg_path and os.path.exists(cfg_path):
        with open(cfg_path, encoding='utf-8') as f:
            return json.load(f)
    return {}


def _results_value(results_dir, spec):
    """spec: {"file": "…", "field": "…"} → 读 json 取字段；field 支持点路径或直接键扫描。"""
    fpath = os.path.join(results_dir, spec['file'])
    if not os.path.exists(fpath):
        return None
    data = json.load(open(fpath, encoding='utf-8'))
    field = spec.get('field')
    if not field:
        return data

    cur = data
    for key in str(field).split('.'):
        if isinstance(cur, dict):
            cur = cur.get(key)
        else:
            cur = None
            break
    if isinstance(cur, (int, float)):
        return float(cur)

    def scan(d):
        if isinstance(d, dict):
            for k, v in d.items():
                if k == field and isinstance(v, (int, float)):
                    return float(v)
                r = scan(v)
                if r is not None:
                    return r
        return None
    return scan(data)


def _csv_rows(results_dir, fname):
    fpath = os.path.join(results_dir, fname)
    rows = []
    if os.path.exists(fpath):
        for line in open(fpath, encoding='utf-8', errors='replace'):
            line = line.strip()
            if not line:
                continue
            rows.append(line.split(','))
    return rows


# ---------------- 各检查 ----------------
def check_placeholder(text):
    issues = []
    for pat, name in [
        (r'\\textbf\{[^}]*句\}', '占位式段首（"…句"）'),
        (r'TODO|XXX|待补|占位|待定|（待|【|TBD|待填写', '待办/编辑残留'),
    ]:
        for m in re.finditer(pat, text):
            issues.append(('C6a ' + name, f'L{_ln(text, m)}: {m.group()[:40]}'))
    return issues


def check_numbering(text, source_type):
    issues = []
    if source_type == 'tex':
        for m in re.finditer(r'\\caption\{(表|图)\s*\d', text):
            issues.append(('C6b caption 手写编号撞自动编号',
                           f'L{_ln(text, m)}: {m.group(0)[:36]}'))
    # Word 与 tex 共同：只查「题注行」重号（行首即 图N/表N），不查正文重复引用
    seen = {}
    for m in re.finditer(r'(?m)^(图|表)\s*(\d+)[:：.、\s]', text):
        key = (m.group(1), int(m.group(2)))
        if key in seen:
            issues.append(('C6b 图/表题注重号',
                           f'L{_ln(text, m)}: 以"{m.group(0).strip()}"开头的题注与 L{seen[key]} 重复'))
        else:
            seen[key] = _ln(text, m)
    return issues


def check_orphans(text, res_nums, tol, config, relax):
    issues = []
    ignore_ctx = config.get('ignore_number_contexts', [])
    ignore_vals = set(float(x) for x in config.get('ignore_numbers', []))
    prose = _plain(text)
    for m in re.finditer(r'(?<![A-Za-z_])\d+\.\d+(?:[eE][+-]?\d+)?', prose):
        v = float(m.group())
        if v in ignore_vals:
            continue
        if any(k in _ctx(prose, m) for k in ignore_ctx):
            continue
        if not _near(v, res_nums, tol):
            issues.append(('C1 孤儿数字（results 无同源）' if not relax else 'C1 孤儿数字（警告）',
                           f'L{_ln(prose, m)}: {v:g}  …{_ctx(prose, m)}…'))
    return issues


def check_anchors(text, config, results_dir, tol):
    issues, coverage = [], []
    plain = _plain(text)
    for a in config.get('anchors', []):
        pat = a['pattern']
        matches = list(re.finditer(pat, plain))
        coverage.append((a.get('name', pat[:20]), len(matches)))
        want = a.get('value')
        if want is None and a.get('file'):
            want = _results_value(results_dir, {'file': a['file'], 'field': a.get('field')})
        if want is None:
            issues.append(('C1 锚点（无法取得真值）', f'{a.get("name")}: 配置缺少 value 或 file 不存在'))
            continue
        atol = float(a.get('tol', tol))
        want = float(want)
        for m in matches:
            mm = re.search(r'(\d+(?:\.\d+)?)', m.group(0))
            if not mm:
                continue
            v = float(mm.group(1))
            if abs(v - want) / max(abs(want), 1e-9) > atol:
                issues.append(('C1 口径锚点与结果不符',
                               f'L{_ln(plain, m)}: {a.get("name")} "{_ctx(plain, m)}" 声称 {v:g}，真值 {want:g}'))
    return issues, coverage


def check_claims(text, config, results_dir):
    issues, coverage = [], []
    plain = _plain(text)
    for c in config.get('claims', []):
        matches = list(re.finditer(c['pattern'], plain))
        coverage.append((c.get('name', c['pattern'][:20]), len(matches)))
        want = c.get('expected')
        if want is None and c.get('expected_file'):
            want = _results_value(results_dir, {'file': c['expected_file'], 'field': c.get('expected_field')})
        if want is None:
            issues.append(('C5 声称检查（无法取得期望值）', f'{c.get("name")}: 配置缺 expected 或 expected_file'))
            continue
        want_nums = {float(x) for x in (want if isinstance(want, list) else [want])}
        for m in matches:
            got = {float(x) for x in re.findall(c.get('extract', r'\d+(?:\.\d+)?'), m.group(0))}
            if not got:
                continue
            if got != want_nums:
                issues.append(('C5 表述与实现不符',
                               f'L{_ln(plain, m)}: {c.get("name")} 声称 {sorted(got)}，'
                               f'实现/结果 {sorted(want_nums)}（{c.get("what", "")}）'))
    return issues, coverage


def check_judgements(text, config, results_dir):
    issues, coverage = [], []
    for j in config.get('judgements', []):
        rows = _csv_rows(results_dir, j['cond_file'])
        if not rows:
            issues.append(('C1① 判据（结果文件缺失）', j.get('name', j['cond_file'])))
            coverage.append((j.get('name', j['cond_file']), 0))
            continue
        ac, cc = int(j.get('angle_col', 0)), int(j.get('cond_col', 1))
        g = {}
        for r in rows:
            if len(r) <= max(ac, cc):
                continue
            try:
                th = float(r[ac])
            except ValueError:
                continue
            v = r[cc]
            g[th] = float('inf') if v.strip().lower() in ('inf', 'nan', '') else float(v)
        thresh = float(j.get('threshold', 1000.0))
        sick = tuple(j.get('sick_words', ['病态', '退化']))
        okw = tuple(j.get('ok_words', ['适定']))
        plain = _plain(text)
        candidates = 0
        for m in re.finditer(r'[^。；\n]+', plain):
            seg = m.group(0)
            ths = [float(x) for x in re.findall(j.get('angle_regex', r'(?:\\theta|θ)\s*[=≈~]\s*(\d+(?:\.\d+)?)'), seg)]
            if not ths or not g:
                continue
            candidates += 1
            is_sick = any(w in seg for w in sick)
            is_ok = any(w in seg for w in okw)
            for th in ths:
                k = min(g, key=lambda t: abs(t - th))
                k = g[k]
                if is_sick and k <= thresh:
                    issues.append(('C1① 病态断言与结果不符',
                                   f'L{_ln(plain, m)}: 断言 θ={th:g}° 病态，'
                                   f'{j["cond_file"]} 该处 κ≈{k:.1f}（≤{thresh:g} 适定）'))
                if is_ok and (k > thresh or k == float('inf')):
                    issues.append(('C1① 适定断言与结果不符',
                                   f'L{_ln(plain, m)}: 断言 θ={th:g}° 适定，'
                                   f'{j["cond_file"]} 该处 κ≈{k:.1f}（>{thresh:g} 病态）'))
        coverage.append((j.get('name', j['cond_file']), candidates))
    return issues, coverage


# ---------------- 汇总 ----------------
def check_all(text, source_type, results_dir, tol, config, relax):
    issues, coverage = [], []
    issues += check_placeholder(text)
    issues += check_numbering(text, source_type)
    issues += check_orphans(text, _load_res_nums(results_dir), tol, config, relax)
    for fn in (check_anchors, check_claims, check_judgements):
        i, c = fn(text, config, results_dir, tol) if fn is check_anchors else fn(text, config, results_dir)
        issues += i
        coverage += c
    return issues, coverage


def report(issues, tag, coverage=None):
    print(f'=== 确定性终审 lint {tag} ===')
    if not issues:
        print('未命中任何缺陷。')
    else:
        seen = {}
        for kind, msg in issues:
            seen.setdefault(kind, []).append(msg)
        for kind, msgs in seen.items():
            print(f'\n[{kind}] × {len(msgs)}')
            for msg in msgs[:10]:
                print('   ', msg)
        print(f'\n共 {len(issues)} 条')
    if coverage is not None:
        if coverage:
            print('\n[配置覆盖] 每个配置检查在正文的实际命中句数：')
            for name, cnt in coverage:
                flag = '  ⚠️ 0 命中（正则未匹配正文，请核对配置或确认已删改）' if cnt == 0 else ''
                print(f'   {name}: {cnt} 处{flag}')
        else:
            print('\n[配置覆盖] 未配置 anchors/claims/judgements，仅运行内置 C6a/C6b/C1 孤儿数字。')


# ---------------- demo：注入已知缺陷，自证命中 ----------------
def run_demo(text, source_type, results_dir, tol, config):
    injected = list(text)
    defects = []   # (缺陷名, 期望捕获的 kind 子串, 注入片段)

    inject = lambda s: injected.append('\n' + s + '\n')

    inject('TODO 待补注入')                       # C6a
    defects.append(('C6a 占位符', 'C6a', 'TODO 待补注入'))

    if source_type == 'tex':
        inject('\\caption{表 9 注入双重编号}')     # C6b tex 手写编号
        defects.append(('C6b caption 手写编号', 'C6b', '\\caption{表 9 注入双重编号}'))
    inject('图 1 注入重号甲\n图 1 注入重号乙')       # C6b Word 重号（tex 也会命中）
    defects.append(('C6b 编号重号', 'C6b', '图 1 注入重号'))

    res_nums = _load_res_nums(results_dir)
    # 孤儿数字注入值必须远离所有 results 数值（含相对容差带）
    base = max(res_nums) + 10.0 if res_nums else 12345.678
    ghost = base
    while _near(ghost, res_nums, tol):
        ghost = ghost * 1.37 + 7.7
    inject(f'本段注入孤儿数字 {ghost:.3f} 以证明命中')
    defects.append(('C1 孤儿数字', 'C1 孤儿数字', f'{ghost:.3f}'))

    if config.get('anchors'):
        inject('初始最大偏差注入 $12.4$ m')
        defects.append(('C1 锚点', 'C1 口径锚点', '注入 $12.4$ m'))
    if config.get('claims'):
        inject('首轮取注入 FY02、FY05')
        defects.append(('C5 声称', 'C5 ', '注入 FY02、FY05'))
    if config.get('judgements'):
        inject('注入误判：$\\theta=180^{\\circ}$ 处近共线为病态')
        defects.append(('C1① 判据', 'C1①', '为病态'))

    injected_text = ''.join(injected)
    issues, coverage = check_all(injected_text, source_type, results_dir, tol, config, relax=False)
    kinds = [k for k, _ in issues]

    ok = True
    print('=== 确定性终审 lint --demo（注入自证）===')
    for name, sub, frag in defects:
        hit = any(sub in k for k in kinds)
        ok &= hit
        print(f'  [{"命中" if hit else "未命中!!"}] {name}  ← 注入「{frag[:30]}」')
    report(issues, '（注入后）', coverage)
    return ok


# ---------------- 入口 ----------------
def main():
    ap = argparse.ArgumentParser(description='论文终审确定性 lint（A 类 C1/C5/C6）')
    ap.add_argument('--paper', required=True, help='main.tex / 完整论文.docx / tex 项目目录')
    ap.add_argument('--results', required=True, help='results 目录（*.json/*.csv）')
    ap.add_argument('--config', default=None, help='项目配置 review_lint.json')
    ap.add_argument('--demo', action='store_true', help='注入已知缺陷自证命中')
    ap.add_argument('--relax-orphans', action='store_true', help='孤儿数字降级为警告（不 FAIL）')
    ap.add_argument('--tol', type=float, default=0.02, help='同源相对容差（默认 0.02）')
    args = ap.parse_args()

    config = load_config(args.config)
    if args.demo:
        text, st = load_text(args.paper)
        ok = run_demo(text, st, args.results, args.tol, config)
        sys.exit(0 if ok else 1)

    text, st = load_text(args.paper)
    issues, coverage = check_all(text, st, args.results, args.tol, config, relax=args.relax_orphans)
    report(issues, '（冻结稿）', coverage)
    fail = bool(issues) or any(cnt == 0 for _, cnt in coverage)
    sys.exit(1 if fail else 0)


if __name__ == '__main__':
    main()
