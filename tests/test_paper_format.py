import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


SCRIPTS = Path(__file__).resolve().parents[1] / "tools" / "docx" / "scripts"
sys.path.insert(0, str(SCRIPTS))

import paper_format as pf


class PaperFormatTests(unittest.TestCase):
    def _front_matter(self):
        doc = pf.new_document()
        pf.title(doc, "题目")
        pf.abstract_title(doc)
        pf.body(doc, "摘要正文")
        pf.keywords(doc, "优化；预测")
        return doc

    def test_reference_template_styles_are_kept_but_sample_body_is_cleared(self):
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "template.docx"
            source = Document()
            source.add_paragraph("模板示例正文，不应进入论文")
            source.save(template)

            doc = pf.new_document(template_path=template)

        self.assertNotIn("模板示例正文", "\n".join(p.text for p in doc.paragraphs))

    def test_official_fixed_template_content_can_be_preserved(self):
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "official.docx"
            source = Document()
            source.add_paragraph("官方固定摘要页")
            source.save(template)

            doc = pf.new_document(
                template_path=template,
                preserve_template_content=True,
            )

        self.assertIn("官方固定摘要页", "\n".join(p.text for p in doc.paragraphs))

    def test_cumcm_structure_validator_requires_abstract_and_keywords(self):
        doc = pf.new_document()
        pf.title(doc, "题目")

        errors = pf.validate_paper_structure(doc)

        self.assertTrue(any("摘要" in error for error in errors))
        self.assertTrue(any("关键词" in error for error in errors))

    def test_complete_cumcm_front_matter_passes(self):
        doc = self._front_matter()

        errors = pf.validate_paper_structure(doc, quality_checks=False)

        self.assertEqual(errors, [])

    def test_quality_validation_reports_length_formula_table_and_page_gaps(self):
        doc = self._front_matter()

        issues = pf.validate_paper_structure(
            doc,
            min_content_units=1200,
            min_equations=2,
            min_tables=1,
            target_pages=10,
        )

        for expected in ("1200", "公式", "表", "渲染页数"):
            self.assertTrue(any(expected in issue for issue in issues), expected)
        self.assertFalse(any("仅检测到 0 幅图" in issue for issue in issues))

    def test_no_cross_competition_minimums_without_explicit_budget(self):
        doc = self._front_matter()

        issues = pf.validate_paper_structure(doc, rendered_pages=1)

        self.assertFalse(any("低于" in issue and "质量目标" in issue for issue in issues))

    def test_table_caption_must_be_referenced_in_body(self):
        doc = self._front_matter()
        pf.body(doc, "正文没有引用下面的表格。")
        pf.three_line_table(doc, [["变量", "值"], ["x", "1"]])
        doc.add_paragraph("表1 参数结果")

        issues = pf.validate_paper_structure(
            doc,
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            require_rendered_pages=False,
        )

        self.assertTrue(any("表1" in issue and "正文" in issue for issue in issues))

    def test_reference_list_and_body_citations_are_bidirectionally_checked(self):
        doc = self._front_matter()
        pf.body(doc, "已有研究支持该方法[1]，但错误引用了[3]。")
        doc.add_paragraph("参考文献")
        doc.add_paragraph("[1] A. Author. A useful paper.")
        doc.add_paragraph("[2] B. Author. An uncited paper.")

        issues = pf.validate_paper_structure(
            doc,
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            require_rendered_pages=False,
        )

        self.assertTrue(any("[3]" in issue and "参考文献表" in issue for issue in issues))
        self.assertTrue(any("[2]" in issue and "未在正文引用" in issue for issue in issues))

    def test_compound_reference_citations_are_recognized(self):
        doc = self._front_matter()
        pf.body(doc, "相关方法见文献[1, 2]及文献[3-4]。")
        doc.add_paragraph("参考文献")
        for number in range(1, 5):
            doc.add_paragraph(f"[{number}] Reference {number}.")

        issues = pf.validate_paper_structure(
            doc,
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            require_rendered_pages=False,
        )

        self.assertFalse(any("未在正文引用" in issue for issue in issues))

    def test_rendered_page_limits_distinguish_target_from_official_maximum(self):
        doc = self._front_matter()

        issues = pf.validate_paper_structure(
            doc,
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            rendered_pages=31,
            official_max_pages=30,
        )

        self.assertTrue(any("官方上限" in issue and "30" in issue for issue in issues))

    def test_safe_save_rejects_skill_root(self):
        doc = self._front_matter()

        with self.assertRaisesRegex(ValueError, "PROJECT_ROOT"):
            pf.save_document(doc, pf.SKILL_ROOT)

    def test_completion_gate_rejects_incomplete_docx(self):
        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "完整论文.docx"
            self._front_matter().save(path)

            report = pf.validate_document(
                path,
                rendered_pages=7,
                min_content_units=1200,
            )

        self.assertFalse(report["passed"])
        self.assertLess(report["metrics"]["content_units"], 1200)
        self.assertTrue(any("1200" in issue for issue in report["issues"]))


if __name__ == "__main__":
    unittest.main()
